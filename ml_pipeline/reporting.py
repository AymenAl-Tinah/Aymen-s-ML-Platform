"""
Report Generation Module
=========================
Creates professional documentation and reports in DOCX format.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
import numpy as np
from typing import Dict, List, Optional
import logging
from datetime import datetime
import io
import base64
import os
from PIL import Image

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    Professional report generation system for ML projects.
    """
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Setup custom styles for the document."""
        styles = self.doc.styles
        
        # Title style
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading style
        if 'CustomHeading' not in styles:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 102, 204)
    
    def add_title(self, title: str, subtitle: str = None):
        """
        Add title and subtitle to the document.
        
        Args:
            title: Main title
            subtitle: Optional subtitle
        """
        # Main title
        title_para = self.doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if subtitle:
            subtitle_para = self.doc.add_paragraph(subtitle)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_para.runs[0].font.size = Pt(14)
            subtitle_para.runs[0].font.italic = True
        
        # Add date
        date_para = self.doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        
        self.doc.add_paragraph()  # Spacing
    
    def add_section(self, title: str, content: str = None):
        """
        Add a new section with heading.
        
        Args:
            title: Section title
            content: Optional section content
        """
        self.doc.add_heading(title, level=1)
        if content:
            self.doc.add_paragraph(content)
    
    def add_subsection(self, title: str, content: str = None):
        """
        Add a subsection with heading.
        
        Args:
            title: Subsection title
            content: Optional subsection content
        """
        self.doc.add_heading(title, level=2)
        if content:
            self.doc.add_paragraph(content)
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """
        Add a paragraph to the document.
        
        Args:
            text: Paragraph text
            bold: Whether to make text bold
            italic: Whether to make text italic
        """
        para = self.doc.add_paragraph(text)
        if bold:
            para.runs[0].font.bold = True
        if italic:
            para.runs[0].font.italic = True
    
    def add_bullet_list(self, items: List[str]):
        """
        Add a bullet list to the document.
        
        Args:
            items: List of items
        """
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')
    
    def add_table(self, df: pd.DataFrame, title: str = None):
        """
        Add a table from a DataFrame.
        
        Args:
            df: DataFrame to convert to table
            title: Optional table title
        """
        if title:
            self.doc.add_paragraph(title, style='Heading 3')
        
        # Create table
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        
        # Add header
        header_cells = table.rows[0].cells
        for idx, column in enumerate(df.columns):
            header_cells[idx].text = str(column)
            header_cells[idx].paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for idx, value in enumerate(row):
                if pd.isna(value):
                    row_cells[idx].text = 'N/A'
                elif isinstance(value, float):
                    row_cells[idx].text = f'{value:.4f}'
                else:
                    row_cells[idx].text = str(value)
        
        self.doc.add_paragraph()  # Spacing
    
    def add_key_value_pairs(self, data: Dict, title: str = None):
        """
        Add key-value pairs as a formatted list.
        
        Args:
            data: Dictionary of key-value pairs
            title: Optional title
        """
        if title:
            self.doc.add_paragraph(title, style='Heading 3')
        
        for key, value in data.items():
            para = self.doc.add_paragraph()
            para.add_run(f"{key}: ").bold = True
            
            if isinstance(value, float):
                para.add_run(f"{value:.4f}")
            elif isinstance(value, (list, dict)):
                para.add_run(str(value))
            else:
                para.add_run(str(value))
    
    def add_dataset_summary(self, dataset_info: Dict):
        """
        Add dataset summary section.
        
        Args:
            dataset_info: Dictionary with dataset information
        """
        self.add_section("Dataset Summary")
        
        summary_data = {
            "Number of Samples": dataset_info.get('n_samples', 'N/A'),
            "Number of Features": dataset_info.get('n_features', 'N/A'),
            "Target Column": dataset_info.get('target_column', 'N/A'),
            "Task Type": dataset_info.get('task_type', 'N/A'),
            "Missing Values": dataset_info.get('missing_values', 'N/A'),
            "Duplicate Rows": dataset_info.get('duplicate_rows', 'N/A')
        }
        
        self.add_key_value_pairs(summary_data)
        
        # Feature types
        if 'numerical_features' in dataset_info:
            self.add_paragraph(f"\nNumerical Features ({len(dataset_info['numerical_features'])}): " + 
                             ", ".join(dataset_info['numerical_features'][:10]))
        
        if 'categorical_features' in dataset_info:
            self.add_paragraph(f"Categorical Features ({len(dataset_info['categorical_features'])}): " + 
                             ", ".join(dataset_info['categorical_features'][:10]))
    
    def add_preprocessing_summary(self, preprocessing_info: Dict):
        """
        Add preprocessing summary section.
        
        Args:
            preprocessing_info: Dictionary with preprocessing information
        """
        self.add_section("Data Preprocessing")
        
        if 'transformations' in preprocessing_info:
            self.add_subsection("Applied Transformations")
            self.add_bullet_list(preprocessing_info['transformations'])
        
        summary_data = {
            "Original Shape": str(preprocessing_info.get('original_shape', 'N/A')),
            "Final Shape": str(preprocessing_info.get('final_shape', 'N/A')),
            "Total Features": preprocessing_info.get('total_features', 'N/A'),
            "Numerical Features": preprocessing_info.get('numerical_features', 'N/A'),
            "Categorical Features": preprocessing_info.get('categorical_features', 'N/A')
        }
        
        self.add_key_value_pairs(summary_data, "Preprocessing Summary")
    
    def add_model_comparison(self, comparison_df: pd.DataFrame):
        """
        Add model comparison table.
        
        Args:
            comparison_df: DataFrame with model comparison
        """
        self.add_section("Model Comparison")
        self.add_table(comparison_df, "Performance Metrics")
    
    def add_best_model_details(self, model_info: Dict):
        """
        Add best model details section.
        
        Args:
            model_info: Dictionary with best model information
        """
        self.add_section("Best Model Details")
        
        details = {
            "Model Name": model_info.get('model_name', 'N/A'),
            "Model Type": model_info.get('model_type', 'N/A'),
            "Training Time": f"{model_info.get('training_time', 0):.2f} seconds"
        }
        
        self.add_key_value_pairs(details)
        
        # Add metrics
        if 'metrics' in model_info:
            self.add_subsection("Performance Metrics")
            self.add_key_value_pairs(model_info['metrics'])
        
        # Add hyperparameters
        if 'hyperparameters' in model_info and model_info['hyperparameters']:
            self.add_subsection("Best Hyperparameters")
            self.add_key_value_pairs(model_info['hyperparameters'])
    
    def add_visualizations(self, visualizations: Dict):
        """
        Add visualizations section with charts.
        
        Args:
            visualizations: Dictionary with plot images (file paths or base64 encoded)
        """
        logger.info(f"=== ADDING VISUALIZATIONS - Total: {len(visualizations)} ===")
        
        plot_titles = {
            'missing_values': 'Missing Values Analysis',
            'distribution': 'Feature Distributions',
            'correlation': 'Correlation Matrix',
            'categorical': 'Categorical Features Analysis',
            'target': 'Target Variable Distribution',
            'outliers': 'Outliers Detection'
        }
        
        added_count = 0
        
        for plot_key, plot_title in plot_titles.items():
            logger.info(f"Processing {plot_key}...")
            if plot_key in visualizations and visualizations[plot_key]:
                try:
                    self.add_subsection(plot_title)
                    
                    plot_data = visualizations[plot_key]
                    logger.info(f"Plot data type: {type(plot_data)}, Value: {plot_data if isinstance(plot_data, str) and len(str(plot_data)) < 100 else 'long string'}")
                    
                    # Check if it's a file path
                    if isinstance(plot_data, str) and os.path.exists(plot_data):
                        # It's a file path - add directly
                        logger.info(f"✓ ADDING IMAGE FROM FILE: {plot_data}")
                        self.doc.add_picture(plot_data, width=Inches(6.0))
                        last_paragraph = self.doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add spacing
                        self.doc.add_paragraph()
                        added_count += 1
                        logger.info(f"✓✓✓ SUCCESS - Added {plot_title} ✓✓✓")
                        
                    elif isinstance(plot_data, str):
                        logger.info(f"Trying base64 decode for {plot_key}")
                        # Try base64 decoding
                        try:
                            if plot_data.startswith('data:image'):
                                plot_data = plot_data.split(',')[1]
                            
                            image_data = base64.b64decode(plot_data)
                            image_stream = io.BytesIO(image_data)
                            
                            # Add image to document
                            self.doc.add_picture(image_stream, width=Inches(6.0))
                            last_paragraph = self.doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Add spacing
                            self.doc.add_paragraph()
                            added_count += 1
                            logger.info(f"✓✓✓ SUCCESS - Added {plot_title} from base64 ✓✓✓")
                        except Exception as decode_error:
                            logger.error(f"XXX Base64 decode FAILED for {plot_key}: {str(decode_error)}")
                            self.add_paragraph(f"[Visualization {plot_title} could not be generated]")
                    else:
                        logger.error(f"XXX Invalid plot data type for {plot_key}: {type(plot_data)}")
                        self.add_paragraph(f"[Visualization {plot_title} could not be generated]")
                        
                except Exception as e:
                    logger.error(f"XXX FAILED to add {plot_key}: {str(e)}")
                    import traceback
                    logger.error(traceback.format_exc())
                    self.add_paragraph(f"[Visualization {plot_title} could not be generated]")
            else:
                logger.warning(f"Skipping {plot_key} - not in visualizations dict or empty")
        
        logger.info(f"=== FINISHED ADDING VISUALIZATIONS - Added {added_count} out of {len(plot_titles)} ===")    
    
    def add_conclusions(self, conclusions: List[str]):
        """
        Add conclusions section.
        
        Args:
            conclusions: List of conclusion points
        """
        self.add_section("Conclusions and Recommendations")
        self.add_bullet_list(conclusions)
    
    def add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()
    
    def save(self, filepath: str):
        """
        Save the document to a file.
        
        Args:
            filepath: Path to save the document
        """
        self.doc.save(filepath)
        logger.info(f"Report saved to: {filepath}")
    
    def generate_complete_report(self, project_data: Dict, filepath: str):
        """
        Generate a complete ML project report.
        
        Args:
            project_data: Dictionary with all project information
            filepath: Path to save the report
        """
        # Title page
        self.add_title(
            "Machine Learning Project Report",
            project_data.get('project_name', 'Automated ML Analysis')
        )
        
        # Executive Summary
        self.add_section("Executive Summary")
        self.add_paragraph(
            f"This report presents a comprehensive analysis of a machine learning project "
            f"for {project_data.get('task_type', 'classification/regression')} task. "
            f"The analysis includes data exploration, preprocessing, model training, "
            f"evaluation, and recommendations."
        )
        
        self.add_page_break()
        
        # Dataset Summary
        if 'dataset_info' in project_data:
            self.add_dataset_summary(project_data['dataset_info'])
            self.add_page_break()
        
        # Preprocessing
        if 'preprocessing_info' in project_data:
            self.add_preprocessing_summary(project_data['preprocessing_info'])
            self.add_page_break()
        
        # Data Visualizations (NEW - in the middle!)
        if project_data.get('include_visualizations'):
            logger.info("=== VISUALIZATIONS REQUESTED ===")
            if 'visualizations' in project_data:
                logger.info(f"Found {len(project_data['visualizations'])} visualizations")
                self.add_section("Data Visualizations")
                self.add_paragraph(
                    "The following visualizations provide insights into the dataset characteristics, "
                    "distributions, correlations, and patterns."
                )
                self.add_visualizations(project_data['visualizations'])
                self.add_page_break()
                logger.info("=== VISUALIZATIONS ADDED TO REPORT ===")
            else:
                logger.error("=== NO VISUALIZATIONS IN PROJECT_DATA ===")
                self.add_section("Data Visualizations")
                self.add_paragraph("[Visualizations could not be generated - please generate visualizations first]")
                self.add_page_break()
        
        # Model Comparison
        if 'comparison_df' in project_data:
            self.add_model_comparison(project_data['comparison_df'])
            self.add_page_break()
        
        # Best Model
        if 'best_model_info' in project_data:
            self.add_best_model_details(project_data['best_model_info'])
            self.add_page_break()
        
        # Conclusions
        conclusions = project_data.get('conclusions', [
            f"Successfully trained and evaluated {project_data.get('n_models', 'multiple')} models",
            f"Best performing model: {project_data.get('best_model_name', 'N/A')}",
            "All models and transformers have been saved for future use",
            "Consider ensemble methods for potentially better performance",
            "Regular model retraining recommended as new data becomes available"
        ])
        self.add_conclusions(conclusions)
        
        # Save report
        self.save(filepath)
        
        logger.info("Complete report generated successfully")
